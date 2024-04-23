import streamlit as st
import tempfile
import os
import pythoncom
import win32com.client
from win32com.client import constants
import time
from docx import Document
from pdf2docx import Converter
import unicodedata
import base64

st.set_page_config(page_title="Document Comparison Tool", page_icon=":memo:", layout="wide")

def main():
    file_01 = st.file_uploader("Upload the first document", type=["pdf", "docx", "doc"])
    file_02 = st.file_uploader("Upload the second document", type=["pdf", "docx", "doc"])
   
    if st.button("Create Full Comparison"):
        if file_01 is not None and file_02 is not None:
            compare_docs(file_01, file_02)
            comparison_path = create_comparison()
            st.write("Documents compared successfully!")
            with open(comparison_path, "rb") as f:
                file = f.read()
                b64 = base64.b64encode(file).decode()
                href = f'<a href="data:file/docx;base64,{b64}" download="comparison.docx">Download Comparison Document</a>'
                st.markdown(href, unsafe_allow_html=True)
            st.write("Click the button below to compare another set of documents.")
            if st.button("Compare Another Set of Documents"):
                st.experimental_rerun()
        else:
            st.write("Please upload both documents before generating the comparison.")
            
    if st.button("Create Summary"):
        if file_01 is not None and file_02 is not None:
            if 'comparison_path' not in st.session_state:
                compare_docs(file_01, file_02)
                st.session_state.comparison_path = create_comparison()
            create_summary(st.session_state.comparison_path)
            st.write("Documents compared successfully!")
            with open(st.session_state.changes_summary_path, "rb") as f:
                file = f.read()
                b64 = base64.b64encode(file).decode()
                href = f'<a href="data:file/docx;base64,{b64}" download="summary.docx">Download Summary Document</a>'
                st.markdown(href, unsafe_allow_html=True)
            if st.button("Compare Another Set of Documents"):
                st.experimental_rerun()
        else:
            st.write("Please upload both documents before generating the summary.")
    return file_01, file_02
        
def compare_docs(file_01, file_02):
    if file_01 is not None and file_02 is not None:
        st.write("Comparing documents...")

        # Save the file paths in session state
        st.session_state.file_01_path = tempfile.mktemp(suffix=".docx")
        st.session_state.file_02_path = tempfile.mktemp(suffix=".docx")
        st.session_state.summary_path = tempfile.mktemp(suffix=".docx")  

        with open(st.session_state.file_01_path, "wb") as f:
            if file_01.name.endswith('.pdf'):
                cv = Converter(file_01.name)
                cv.convert(st.session_state.file_01_path, start=0, end=None)
                cv.close()
            elif file_01.name.endswith('.doc'):
                doc = Document(file_01)
                doc.save(st.session_state.file_01_path)
            else:
                f.write(file_01.getbuffer())

        with open(st.session_state.file_02_path, "wb") as f:
            if file_02.name.endswith('.pdf'):
                cv = Converter(file_02.name)
                cv.convert(st.session_state.file_02_path, start=0, end=None)
                cv.close()
            elif file_02.name.endswith('.doc'):
                doc = Document(file_02)
                doc.save(st.session_state.file_02_path)
            else:
                f.write(file_02.getbuffer())


def create_comparison():
    # Initialize the COM library
    pythoncom.CoInitialize()

    Application = win32com.client.gencache.EnsureDispatch("Word.Application")

    file_01 = Application.Documents.Open(os.path.abspath(st.session_state.file_01_path))
    file_02 = Application.Documents.Open(os.path.abspath(st.session_state.file_02_path))


    result = Application.CompareDocuments(file_01, file_02, CompareFormatting=False, CompareCaseChanges=False, CompareWhitespace=False, CompareComments=False)
    time.sleep(1)

    st.session_state.comparison_path = tempfile.mktemp(suffix=".docx")
    for _ in range(10):  # Retry up to 10 times
        try:
            result.SaveAs(FileName=os.path.abspath(st.session_state.comparison_path))
            break  # If the call succeeds, break the loop
        except Exception as e:
            print(f"SaveAs failed with error: {e}, retrying...")
            time.sleep(1)  
    file_01.Close()
    file_02.Close()
    Application.ActiveDocument.Close()

    Application.Quit()
    Application = None
    return st.session_state.comparison_path

            
def create_summary(comparison_path):
    pythoncom.CoInitialize()

    changes_doc = Document()
    changes = []
    progress_bar = st.progress(0)
    Application = win32com.client.Dispatch("Word.Application")
    
    # Open the comparison document and set it as the active document
    comparison_doc = Application.Documents.Open(os.path.abspath(st.session_state.comparison_path))
    Application.ActiveWindow.View.Type = constants.wdPrintView
    Application.ActiveWindow.ActivePane.View.SeekView = constants.wdSeekMainDocument

    revisions = list(Application.ActiveDocument.Revisions)
    rev_count = len(revisions)

    for i, rev in enumerate(revisions, start=1):
        # Get the page number
        page_num = rev.Range.Information(constants.wdActiveEndAdjustedPageNumber)
        # Remove control characters
        text = "".join(ch for ch in rev.Range.Text if unicodedata.category(ch)[0] != "C")
        # Determine the type of the revision
        if rev.Type == constants.wdRevisionInsert:
            revision_type = "Inserted"
        elif rev.Type == constants.wdRevisionDelete:
            revision_type = "Deleted"
        else:
            continue
        # Add the page number, revision type, and the revised text to the changes list
        changes.append(f"Page {page_num}, {revision_type}:\n{text}")
    
        # Update the progress bar
        progress = min((i + 1) / rev_count, 1.0)
        progress_bar.progress(progress)
    # Add the collected changes to the summary document
    for change in changes:
        changes_doc.add_paragraph(change)

    # Save the summary document
    st.session_state.changes_summary_path = tempfile.mktemp(suffix=".docx")
    changes_doc.save(st.session_state.changes_summary_path)
    comparison_doc.Close()

    Application.Quit()
    Application = None
    return st.session_state.changes_summary_path

main()